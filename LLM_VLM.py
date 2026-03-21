import os
import re
import json
import base64
import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

import requests
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Prefer stdlib zoneinfo (Py3.9+); fall back to naive if unavailable
try:
    from zoneinfo import ZoneInfo
    KSA_TZ = ZoneInfo("Asia/Riyadh")
except Exception:
    KSA_TZ = None

# ==============================
# -------- Config -------------
# ==============================
def _openrouter_chat_url() -> str:
    """POST-only URL; a GET to this path returns 404 from OpenRouter."""
    base = os.environ.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1").strip().rstrip("/")
    return f"{base}/chat/completions"


def _openrouter_vlm_model() -> str:
    # Default: vision-capable model known to exist on OpenRouter (InternVL3 slug was removed).
    return os.environ.get("OPENROUTER_VLM_MODEL", "openai/gpt-4o-mini").strip()

# Defaults (can be overridden by CLI)
IMAGE_PATH  = "/content/site_photo.jpeg"
OUT_DIR     = "/content"
RAW_TXT     = "hse_raw.txt"
JSON_PATH   = "hse_report.json"
DOCX_PATH   = "hse_report.docx"
MAX_RECORDS = 5

# ==============================
# ---- OpenRouter helpers ------
# ==============================
def _get_api_key() -> str:
    key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not key:
        raise RuntimeError("Missing OPENROUTER_API_KEY env var.")
    return key

def _image_to_data_url(path: str) -> str:
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    ext = Path(path).suffix.lower()
    mime = "image/png" if ext == ".png" else "image/jpeg"
    return f"data:{mime};base64,{b64}"

def call_openrouter_mm(model: str, prompt: str, image_path: str,
                       max_tokens: int = 1800, temperature: float = 0.0) -> str:
    referer = os.environ.get("OPENROUTER_HTTP_REFERER", "https://localhost").strip()
    title = os.environ.get("OPENROUTER_APP_TITLE", "AUHSE HSE Report Generator").strip()
    headers = {
        "Authorization": f"Bearer {_get_api_key()}",
        "Content-Type": "application/json",
        "HTTP-Referer": referer or "https://localhost",
        "X-Title": title or "AUHSE",
    }
    payload = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": _image_to_data_url(image_path)}}
            ]
        }],
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    url = _openrouter_chat_url()
    r = requests.post(url, headers=headers, json=payload, timeout=120)
    if not r.ok:
        snippet = (r.text or "")[:1500]
        raise RuntimeError(
            f"OpenRouter request failed ({r.status_code}) for {url}: {snippet or r.reason}",
        )
    data = r.json()
    try:
        return data["choices"][0]["message"]["content"]
    except Exception:
        raise RuntimeError(f"Unexpected OpenRouter response: {data}")

# ==============================
# ---- JSON sanitization -------
# ==============================
def _strip_code_fences(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^\s*```(?:json)?\s*", "", s, flags=re.I)
    s = re.sub(r"\s*```\s*$", "", s)
    return s

def _normalize_quotes(s: str) -> str:
    return (s.replace("“", '"').replace("”", '"')
             .replace("‘", "'").replace("’", "'"))

def _remove_comments(s: str) -> str:
    s = re.sub(r"//.*?$", "", s, flags=re.M)
    s = re.sub(r"/\*.*?\*/", "", s, flags=re.S)
    return s

def _remove_trailing_commas(s: str) -> str:
    return re.sub(r",\s*([\]\}])", r"\1", s)

def _sanitize_to_jsonish(text: str) -> str:
    s = _strip_code_fences(text)
    s = _normalize_quotes(s)
    s = _remove_comments(s)
    s = _remove_trailing_commas(s)
    m = re.search(r"(\{.*\}|\[.*\])", s, flags=re.S)
    if m:
        s = m.group(1)
    return s.strip()

def extract_json_lossy(text: str, raw_path: str) -> Dict[str, Any]:
    Path(raw_path).write_text(text, encoding="utf-8")
    try:
        return json.loads(text)
    except Exception:
        pass
    s = _sanitize_to_jsonish(text)
    try:
        return json.loads(s)
    except Exception:
        s2 = re.sub(r"}\s*{", "},\n{", s)
        return json.loads(s2)

# ==============================
# ---- Prompt builder ----------
# ==============================
def _now_ksa():
    if KSA_TZ:
        return datetime.now(KSA_TZ)
    return datetime.now()

def build_hse_prompt(
    image_filename: str,
    project_name: str,
    site_location: str,
    inspection_by: str,
    verified_by: str,
    max_records: int = 5
) -> str:
    """
    Strict JSON-only schema matching the Middle East-style HSE template.
    Names come from user; dates auto-filled (Asia/Riyadh).
    """
    now = _now_ksa()
    date_time   = now.strftime("%Y-%m-%d %H:%M")  # for Project Information
    verify_date = now.strftime("%Y-%m-%d")        # for Follow-up / Verification

    schema = f"""
Return VALID JSON ONLY. No markdown, no code fences, no comments, no trailing commas.
Use double quotes for all keys and string values.

Schema EXACTLY:

{{
  "project_information": {{
    "project_name": "string",
    "site_location": "string",
    "date_time": "YYYY-MM-DD HH:MM",
    "inspection_by": "string"
  }},
  "inspection_summary": {{
    "area_or_zone_inspected": "string",
    "weather_or_environmental_conditions": "string",
    "scope_of_inspection": "string"
  }},
  "detected_hazards": [
    {{
      "image_reference": 1,
      "description_of_hazards": "string",
      "causes": "string",
      "location_on_site": "string"
    }}
  ],
  "risk_analysis": {{
    "likelihood_1_to_5": 1,
    "severity_1_to_5": 1,
    "risk_rating_lxS": 1,
    "risk_level": "Low|Medium|High|Critical"
  }},
  "control_measures": {{
    "elimination": "string",
    "substitution": "string",
    "engineering_controls": "string",
    "administrative_controls": "string",
    "personal_protective_equipment": "string"
  }},
  "recommendations": {{
    "immediate_actions": "string",
    "short_term_measures": "string",
    "long_term_measures": "string"
  }},
  "responsible_and_deadline": {{
    "responsible_department_or_person": "string",
    "due_date": "YYYY-MM-DD"
  }},
  "follow_up_verification": {{
    "verification_date": "YYYY-MM-DD",
    "verified_by": "string",
    "comments": "string"
  }}
}}

MUST-USE CONSTANTS (copy EXACTLY into the JSON; do not alter):
- project_information.project_name = "{project_name}"
- project_information.site_location = "{site_location}"
- project_information.date_time = "{date_time}"
- project_information.inspection_by = "{inspection_by}"
- follow_up_verification.verification_date = "{verify_date}"
- follow_up_verification.verified_by = "{verified_by}"

Instructions:
- Analyze the image THOROUGHLY and provide DETAILED, SPECIFIC, and ACTIONABLE information.
- Use ONLY what can be inferred from the image for the rest of the fields.
- Provide between 3 and {max_records} items in "detected_hazards". Set "image_reference"=1 for all.
- For each hazard, provide COMPREHENSIVE descriptions including:
  * Specific details about what is visible (equipment, structures, people, conditions)
  * Clear identification of safety violations or risks
  * Precise location descriptions (e.g., "northwest corner", "scaffold level 3", "near crane base")
- For causes, explain ROOT CAUSES, not just symptoms. Be specific about why the hazard exists.
- For control measures, provide DETAILED, PRACTICAL, and IMPLEMENTABLE solutions:
  * Elimination: Specific actions to completely remove the hazard
  * Substitution: Alternative methods or materials that reduce risk
  * Engineering Controls: Technical solutions (barriers, guards, ventilation, etc.)
  * Administrative Controls: Procedures, training, signage, work schedules
  * PPE: Specific equipment needed (hard hats, safety vests, gloves, eye protection, etc.)
- For recommendations, provide TIMELINE-SPECIFIC actions:
  * Immediate Actions: What must be done RIGHT NOW (within hours)
  * Short-Term Measures: Actions for the next few days/weeks
  * Long-Term Measures: Strategic improvements for weeks/months ahead
- Compute risk_rating_lxS = likelihood_1_to_5 * severity_1_to_5; choose a consistent risk_level based on:
  * Likelihood: How probable is an incident? (1=Very Unlikely, 5=Almost Certain)
  * Severity: How severe would consequences be? (1=Minor, 5=Catastrophic)
  * Risk Level: Low (1-6), Medium (7-12), High (13-20), Critical (21-25)
- For responsible party, suggest SPECIFIC roles (e.g., "Site Safety Manager", "Project Engineer", "Foreman")
- For due dates, provide REALISTIC timelines based on urgency
- KSA context (SBC standards, Civil Defense regulations, HRSD OSH requirements) should influence practical actions, but DO NOT name them explicitly.
- Be PROFESSIONAL, TECHNICAL, and DETAILED. Avoid generic statements. Provide actionable, site-specific information.

Now analyze the attached image "{image_filename}" in detail and return JSON ONLY matching the schema above with comprehensive, detailed information.
"""
    return schema.strip()

# ==============================
# ---- DOCX rendering ----------
# ==============================
def set_cell_border(cell, **kwargs):
    """Set borders for a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_val = kwargs.get(border_name)
        if border_val:
            tag = f'{border_name}'
            element = OxmlElement(f'w:{tag}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)

def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def write_hse_docx(hse: Dict[str, Any], out_path: str):
    d = Document()
    
    # Set default style
    style = d.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== TITLE PAGE ==========
    # Title
    title_para = d.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('HEALTH, SAFETY & ENVIRONMENT (HSE)\nINSPECTION REPORT')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    d.add_paragraph()  # Spacing
    
    # Subtitle
    subtitle_para = d.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('Middle East Compliance Format')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    
    d.add_paragraph()  # Spacing
    d.add_paragraph()  # Spacing
    
    # ========== 1. PROJECT INFORMATION ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('1. PROJECT INFORMATION')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    pi = hse.get("project_information", {})
    info_table = d.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Set column widths
    info_table.columns[0].width = Inches(2.5)
    info_table.columns[1].width = Inches(4.5)
    
    # Project Name
    row = info_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Project Name:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(pi.get("project_name", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Site Location
    row = info_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Site Location:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(pi.get("site_location", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Date/Time
    row = info_table.rows[2]
    row.cells[0].paragraphs[0].add_run('Date / Time:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(pi.get("date_time", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Inspection By
    row = info_table.rows[3]
    row.cells[0].paragraphs[0].add_run('Inspection By:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(pi.get("inspection_by", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    d.add_paragraph()  # Spacing
    
    # ========== 2. INSPECTION SUMMARY ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('2. INSPECTION SUMMARY')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    ins = hse.get("inspection_summary", {})
    summary_table = d.add_table(rows=3, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_table.columns[0].width = Inches(2.5)
    summary_table.columns[1].width = Inches(4.5)
    
    # Area/Zone Inspected
    row = summary_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Area/Zone Inspected:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(ins.get("area_or_zone_inspected", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Weather/Environmental Conditions
    row = summary_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Weather/Environmental Conditions:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(ins.get("weather_or_environmental_conditions", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Scope of Inspection
    row = summary_table.rows[2]
    row.cells[0].paragraphs[0].add_run('Scope of Inspection:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    scope_text = ins.get("scope_of_inspection", "N/A")
    row.cells[1].paragraphs[0].add_run(scope_text)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    d.add_paragraph()  # Spacing
    
    # ========== 3. DETECTED HAZARDS ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('3. DETECTED HAZARDS')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    hazards = hse.get("detected_hazards", [])
    if hazards:
        for idx, hz in enumerate(hazards, 1):
            # Hazard header
            hazard_para = d.add_paragraph()
            hazard_run = hazard_para.add_run(f'Hazard {idx} (Image Reference: #{hz.get("image_reference", 1)})')
            hazard_run.font.size = Pt(13)
            hazard_run.font.bold = True
            hazard_run.font.color.rgb = RGBColor(192, 0, 0)
            
            # Hazard details table
            hazard_table = d.add_table(rows=3, cols=2)
            hazard_table.style = 'Light Grid Accent 1'
            hazard_table.columns[0].width = Inches(2.5)
            hazard_table.columns[1].width = Inches(4.5)
            
            # Description
            row = hazard_table.rows[0]
            row.cells[0].paragraphs[0].add_run('Description:').bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            desc_text = hz.get("description_of_hazards", "N/A")
            row.cells[1].paragraphs[0].add_run(desc_text)
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            # Causes
            row = hazard_table.rows[1]
            row.cells[0].paragraphs[0].add_run('Causes:').bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            causes_text = hz.get("causes", "N/A")
            row.cells[1].paragraphs[0].add_run(causes_text)
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            # Location
            row = hazard_table.rows[2]
            row.cells[0].paragraphs[0].add_run('Location on Site:').bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            location_text = hz.get("location_on_site", "N/A")
            row.cells[1].paragraphs[0].add_run(location_text)
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            d.add_paragraph()  # Spacing between hazards
    else:
        no_hazards_para = d.add_paragraph()
        no_hazards_para.add_run('No hazards detected during this inspection.').italic = True
    
    d.add_paragraph()  # Spacing
    
    # ========== 4. RISK ANALYSIS ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('4. RISK ANALYSIS')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    ra = hse.get("risk_analysis", {})
    risk_table = d.add_table(rows=4, cols=2)
    risk_table.style = 'Light Grid Accent 1'
    
    risk_table.columns[0].width = Inches(2.5)
    risk_table.columns[1].width = Inches(4.5)
    
    # Likelihood
    row = risk_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Likelihood (1-5):').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    likelihood = ra.get("likelihood_1_to_5", "N/A")
    row.cells[1].paragraphs[0].add_run(str(likelihood))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Severity
    row = risk_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Severity (1-5):').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    severity = ra.get("severity_1_to_5", "N/A")
    row.cells[1].paragraphs[0].add_run(str(severity))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Risk Rating
    row = risk_table.rows[2]
    row.cells[0].paragraphs[0].add_run('Risk Rating (L×S):').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    risk_rating = ra.get("risk_rating_lxS", "N/A")
    row.cells[1].paragraphs[0].add_run(str(risk_rating))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Risk Level with color coding
    row = risk_table.rows[3]
    row.cells[0].paragraphs[0].add_run('Risk Level:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    risk_level = ra.get("risk_level", "N/A")
    risk_run = row.cells[1].paragraphs[0].add_run(str(risk_level).upper())
    risk_run.font.size = Pt(12)
    risk_run.font.bold = True
    
    # Color code based on risk level
    risk_level_lower = str(risk_level).lower()
    if 'critical' in risk_level_lower or 'high' in risk_level_lower:
        risk_run.font.color.rgb = RGBColor(192, 0, 0)  # Red
    elif 'medium' in risk_level_lower:
        risk_run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
    else:
        risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    d.add_paragraph()  # Spacing
    
    # ========== 5. CONTROL MEASURES ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('5. CONTROL MEASURES (Hierarchy of Controls)')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    cm = hse.get("control_measures", {})
    controls_table = d.add_table(rows=5, cols=2)
    controls_table.style = 'Light Grid Accent 1'
    
    controls_table.columns[0].width = Inches(2.8)
    controls_table.columns[1].width = Inches(4.2)
    
    control_items = [
        ("Elimination", cm.get("elimination", "N/A")),
        ("Substitution", cm.get("substitution", "N/A")),
        ("Engineering Controls", cm.get("engineering_controls", "N/A")),
        ("Administrative Controls", cm.get("administrative_controls", "N/A")),
        ("Personal Protective Equipment (PPE)", cm.get("personal_protective_equipment", "N/A"))
    ]
    
    for idx, (label, value) in enumerate(control_items):
        row = controls_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(f'{label}:').bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        row.cells[1].paragraphs[0].add_run(str(value))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    d.add_paragraph()  # Spacing
    
    # ========== 6. RECOMMENDATIONS ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('6. RECOMMENDATIONS')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    rec = hse.get("recommendations", {})
    rec_table = d.add_table(rows=3, cols=2)
    rec_table.style = 'Light Grid Accent 1'
    
    rec_table.columns[0].width = Inches(2.5)
    rec_table.columns[1].width = Inches(4.5)
    
    # Immediate Actions
    row = rec_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Immediate Actions:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(rec.get("immediate_actions", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Short-Term Measures
    row = rec_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Short-Term Measures:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(rec.get("short_term_measures", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Long-Term Measures
    row = rec_table.rows[2]
    row.cells[0].paragraphs[0].add_run('Long-Term Measures:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(rec.get("long_term_measures", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    d.add_paragraph()  # Spacing
    
    # ========== 7. RESPONSIBLE PARTY & DEADLINE ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('7. RESPONSIBLE PARTY & DEADLINE')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    rd = hse.get("responsible_and_deadline", {})
    resp_table = d.add_table(rows=2, cols=2)
    resp_table.style = 'Light Grid Accent 1'
    
    resp_table.columns[0].width = Inches(2.5)
    resp_table.columns[1].width = Inches(4.5)
    
    # Responsible Department/Person
    row = resp_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Responsible Department/Person:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(rd.get("responsible_department_or_person", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Due Date
    row = resp_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Due Date:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    due_date = rd.get("due_date", "N/A")
    due_run = row.cells[1].paragraphs[0].add_run(str(due_date))
    due_run.font.size = Pt(11)
    due_run.font.bold = True
    due_run.font.color.rgb = RGBColor(192, 0, 0)  # Red for urgency
    
    d.add_paragraph()  # Spacing
    
    # ========== 8. FOLLOW-UP / VERIFICATION ==========
    section_para = d.add_paragraph()
    section_run = section_para.add_run('8. FOLLOW-UP / VERIFICATION')
    section_run.font.size = Pt(16)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 51, 102)
    
    fu = hse.get("follow_up_verification", {})
    verify_table = d.add_table(rows=3, cols=2)
    verify_table.style = 'Light Grid Accent 1'
    
    verify_table.columns[0].width = Inches(2.5)
    verify_table.columns[1].width = Inches(4.5)
    
    # Verification Date
    row = verify_table.rows[0]
    row.cells[0].paragraphs[0].add_run('Verification Date:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(fu.get("verification_date", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Verified By
    row = verify_table.rows[1]
    row.cells[0].paragraphs[0].add_run('Verified By:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(fu.get("verified_by", "N/A"))
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Comments
    row = verify_table.rows[2]
    row.cells[0].paragraphs[0].add_run('Comments:').bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    comments_text = fu.get("comments", "N/A")
    row.cells[1].paragraphs[0].add_run(comments_text)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    d.add_paragraph()  # Spacing
    d.add_paragraph()  # Spacing
    
    # Footer
    footer_para = d.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('--- End of Report ---')
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    d.save(out_path)

# ==============================
# ---- Public API --------------
# ==============================
def generate_hse_report(
    image_path: str,
    project_name: str,
    site_location: str,
    inspection_by: str,
    verified_by: str,
    out_dir: str,
    max_records: int = MAX_RECORDS,
) -> Dict[str, Any]:
    """
    Run the OpenRouter vision-LM HSE report pipeline and persist outputs.
    Model: OPENROUTER_VLM_MODEL or default openai/gpt-4o-mini.
    """
    out_dir_path = Path(out_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)

    raw_path = out_dir_path / RAW_TXT
    json_path = out_dir_path / JSON_PATH
    docx_path = out_dir_path / DOCX_PATH

    prompt = build_hse_prompt(
        image_filename=Path(image_path).name,
        project_name=project_name,
        site_location=site_location,
        inspection_by=inspection_by,
        verified_by=verified_by,
        max_records=max_records,
    )

    raw = call_openrouter_mm(
        model=_openrouter_vlm_model(),
        prompt=prompt,
        image_path=image_path,
        max_tokens=1800,
        temperature=0.0,
    )

    obj = extract_json_lossy(raw, raw_path=str(raw_path))
    if not isinstance(obj, dict):
        raise RuntimeError("Expected a JSON object as top-level output.")

    now = _now_ksa()
    obj.setdefault("project_information", {})
    obj["project_information"]["project_name"] = project_name
    obj["project_information"]["site_location"] = site_location
    obj["project_information"]["inspection_by"] = inspection_by
    obj["project_information"]["date_time"] = now.strftime("%Y-%m-%d %H:%M")

    obj.setdefault("follow_up_verification", {})
    obj["follow_up_verification"]["verification_date"] = now.strftime("%Y-%m-%d")
    obj["follow_up_verification"]["verified_by"] = verified_by
    obj["follow_up_verification"].setdefault("comments", "")

    Path(json_path).write_text(
        json.dumps(obj, ensure_ascii=False, indent=2),
        encoding="utf-8-sig",
    )

    write_hse_docx(obj, str(docx_path))

    return {
        "hse": obj,
        "raw_model_output": raw,
        "raw_text_path": str(raw_path),
        "json_path": str(json_path),
        "docx_path": str(docx_path),
        "prompt": prompt,
        "generated_at": now.isoformat(),
    }

# ==============================
# ------------ Main ------------
# ==============================
def main():
    ap = argparse.ArgumentParser(description="Generate Middle East-style HSE Report from image via InternVL3-78B")
    ap.add_argument("--image", default=IMAGE_PATH, help="Input image path")
    ap.add_argument("--out-dir", default=OUT_DIR, help="Output directory")
    # User-provided names
    ap.add_argument("--project-name", required=True)
    ap.add_argument("--site-location", required=True)
    ap.add_argument("--inspection-by", required=True)
    ap.add_argument("--verified-by", required=True)
    ap.add_argument("--max-records", type=int, default=MAX_RECORDS)
    args = ap.parse_args()

    print("[*] Running HSE report generation pipeline...")
    result = generate_hse_report(
        image_path=args.image,
        project_name=args.project_name,
        site_location=args.site_location,
        inspection_by=args.inspection_by,
        verified_by=args.verified_by,
        out_dir=args.out_dir,
        max_records=args.max_records,
    )
    print(f"[OK] JSON saved: {result['json_path']}")
    print(f"[OK] DOCX saved: {result['docx_path']}")
    print("[*] Done.")

if __name__ == "__main__":
    main()
