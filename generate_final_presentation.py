#!/usr/bin/env python3
"""
Generate Final Presentation Document for AUHSE HSE Inspection System
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('University of Hail\nCollege of Computer Science and Engineering\n\n')
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_para2.add_run('Final Presentation\n\n')
    title_run2.font.size = Pt(18)
    title_run2.font.bold = True
    
    title_para3 = doc.add_paragraph()
    title_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run3 = title_para3.add_run('Smart Digital System for Occupational Health, Safety, Environment and Sustainability (HSE&S)\n\nAUHSE: Autonomous HSE Intelligence System\n\n')
    title_run3.font.size = Pt(16)
    title_run3.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    supervisor_para = doc.add_paragraph()
    supervisor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor_para.add_run('Under Supervision of Academic Advisor:\nProf/Dr. [Supervisor Name]\n\n')
    
    student_para = doc.add_paragraph()
    student_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student_para.add_run('Submitted by:\n[Student Name]\n[Student ID]')
    
    doc.add_page_break()
    
    # Table of Contents
    toc_para = doc.add_paragraph()
    toc_para.add_run('Table of Contents').bold = True
    toc_para.add_run().add_break()
    toc_para.add_run('1. Feasibility Analysis')
    toc_para.add_run().add_break()
    toc_para.add_run('2. Requirements Analysis')
    toc_para.add_run().add_break()
    toc_para.add_run('3. System Modeling')
    toc_para.add_run().add_break()
    toc_para.add_run('   3.1 Use Cases')
    toc_para.add_run().add_break()
    toc_para.add_run('   3.2 Class Diagram')
    toc_para.add_run().add_break()
    toc_para.add_run('   3.3 Sequence Diagrams')
    toc_para.add_run().add_break()
    toc_para.add_run('4. Database Design')
    toc_para.add_run().add_break()
    toc_para.add_run('5. Architectural Design')
    toc_para.add_run().add_break()
    toc_para.add_run('6. User Interface Design')
    toc_para.add_run().add_break()
    toc_para.add_run('7. Future Work')
    toc_para.add_run().add_break()
    toc_para.add_run('8. References')
    
    doc.add_page_break()
    
    # 1. FEASIBILITY ANALYSIS
    heading1 = doc.add_heading('1. Feasibility Analysis', level=1)
    
    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph(
        'The AUHSE (Autonomous HSE Intelligence System) is a smart digital system designed to automate '
        'Health, Safety, Environment, and Sustainability inspections for construction sites. This feasibility '
        'analysis evaluates the technical, economic, operational, and schedule feasibility of implementing '
        'this AI-powered inspection system.'
    )
    
    doc.add_heading('1.2 Technical Feasibility', level=2)
    doc.add_paragraph(
        'The system leverages cutting-edge technologies that are proven and readily available:'
    )
    
    doc.add_paragraph('• Computer Vision: YOLOv8 (You Only Look Once) object detection model for real-time '
                     'detection of PPE, workers, and construction equipment. YOLOv8 is a state-of-the-art '
                     'deep learning model optimized for speed and accuracy.', style='List Bullet')
    
    doc.add_paragraph('• Vision-Language Models: InternVL3-78B, a large-scale multimodal model capable of '
                     'understanding both images and text, enabling detailed hazard analysis and report generation.', 
                     style='List Bullet')
    
    doc.add_paragraph('• Web Framework: Flask, a lightweight and flexible Python web framework suitable for '
                     'rapid development and deployment.', style='List Bullet')
    
    doc.add_paragraph('• Cloud APIs: OpenRouter API provides access to InternVL3-78B without requiring local '
                     'GPU infrastructure, reducing hardware costs.', style='List Bullet')
    
    doc.add_paragraph(
        'All required technologies are mature, well-documented, and have active community support. The system '
        'requires minimal hardware infrastructure as it leverages cloud-based AI services.'
    )
    
    doc.add_heading('1.2.1 Risk Regarding Familiarity with Technology', level=3)
    doc.add_paragraph(
        'The development team requires expertise in Python programming, Flask web development, computer vision, '
        'and API integration. These are standard skills in modern software development. The use of pre-trained '
        'models (YOLOv8, InternVL3) significantly reduces the complexity compared to training models from scratch.'
    )
    
    doc.add_heading('1.2.2 Risk Regarding Project Size', level=3)
    doc.add_paragraph(
        'The project scope is well-defined and manageable. The core functionality involves: (1) Image upload and '
        'validation, (2) YOLO-based relevance gate, (3) Vision-language model analysis, and (4) Report generation. '
        'The modular architecture allows for incremental development and testing.'
    )
    
    doc.add_heading('1.3 Economic Feasibility', level=2)
    doc.add_paragraph('Cost-Benefit Analysis:')
    
    doc.add_paragraph('Development Costs:', style='List Bullet')
    doc.add_paragraph('   • Software development time: Estimated 3-4 months for a team of 2-3 developers', 
                     style='List Bullet 2')
    doc.add_paragraph('   • Cloud API costs: OpenRouter API usage for InternVL3 (~$0.01-0.05 per image analysis)', 
                     style='List Bullet 2')
    doc.add_paragraph('   • YOLOv8 model: Free and open-source', style='List Bullet 2')
    doc.add_paragraph('   • Infrastructure: Minimal (web server hosting costs ~$10-50/month)', 
                     style='List Bullet 2')
    
    doc.add_paragraph('Benefits:', style='List Bullet')
    doc.add_paragraph('   • Reduced manual inspection time: 80-90% time savings per inspection', 
                     style='List Bullet 2')
    doc.add_paragraph('   • Improved accuracy: Consistent hazard detection and risk assessment', 
                     style='List Bullet 2')
    doc.add_paragraph('   • 24/7 availability: Automated system available anytime', 
                     style='List Bullet 2')
    doc.add_paragraph('   • Standardized reports: Ensures compliance with Middle-East HSE regulations', 
                     style='List Bullet 2')
    doc.add_paragraph('   • Cost savings: Reduced need for multiple on-site HSE inspectors', 
                     style='List Bullet 2')
    
    doc.add_paragraph(
        'The system offers significant return on investment by automating repetitive inspection tasks, '
        'reducing human error, and ensuring regulatory compliance.'
    )
    
    doc.add_heading('1.4 Organizational Feasibility', level=2)
    doc.add_paragraph(
        'The system aligns with modern construction industry trends toward digitalization and automation. '
        'Construction companies are increasingly adopting AI and computer vision technologies for safety '
        'monitoring. The Middle-East construction sector, particularly in Saudi Arabia, is actively pursuing '
        'digital transformation initiatives that align with Vision 2030 goals.'
    )
    
    doc.add_paragraph(
        'The system supports existing HSE workflows by automating report generation while maintaining the '
        'structured format required by regulatory bodies. Human inspectors remain in the loop for verification '
        'and decision-making, ensuring organizational acceptance.'
    )
    
    doc.add_heading('1.5 Schedule Feasibility', level=2)
    doc.add_paragraph('The project can be completed within a typical academic timeline:')
    
    doc.add_paragraph('Phase 1 (Weeks 1-4): Requirements gathering, system design, and architecture planning', 
                     style='List Bullet')
    doc.add_paragraph('Phase 2 (Weeks 5-8): Core development - Flask application, YOLO integration, API setup', 
                     style='List Bullet')
    doc.add_paragraph('Phase 3 (Weeks 9-12): Vision-language model integration, report generation, UI development', 
                     style='List Bullet')
    doc.add_paragraph('Phase 4 (Weeks 13-16): Testing, refinement, documentation, and deployment', 
                     style='List Bullet')
    
    doc.add_paragraph(
        'The use of existing frameworks and APIs significantly reduces development time compared to building '
        'from scratch. The modular design allows for parallel development of different components.'
    )
    
    doc.add_page_break()
    
    # 2. REQUIREMENTS ANALYSIS
    heading2 = doc.add_heading('2. Requirements Analysis', level=1)
    
    doc.add_heading('2.1 Introduction', level=2)
    doc.add_paragraph(
        'Requirements analysis involves identifying, documenting, and validating the functional and '
        'non-functional requirements of the AUHSE system. This section covers the requirements gathered '
        'through analysis of the construction industry needs, HSE compliance standards, and system objectives.'
    )
    
    doc.add_heading('2.2 Functional Requirements', level=2)
    
    doc.add_paragraph('FR1: Image Upload and Validation', style='List Bullet')
    doc.add_paragraph('   The system shall accept image uploads in JPEG and PNG formats up to 16 MB. '
                     'The system shall validate file format and size before processing.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('FR2: Relevance Detection', style='List Bullet')
    doc.add_paragraph('   The system shall use YOLOv8 to detect construction-related objects including '
                     'PPE (helmets, vests, gloves, harnesses), workers, and construction equipment. '
                     'The system shall only process images that contain relevant construction activity.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('FR3: Hazard Detection and Analysis', style='List Bullet')
    doc.add_paragraph('   The system shall analyze uploaded images using InternVL3-78B to identify '
                     'safety hazards, assess risk levels, and determine appropriate control measures.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('FR4: Risk Assessment', style='List Bullet')
    doc.add_paragraph('   The system shall calculate risk ratings using the Likelihood × Severity (L×S) '
                     'matrix, where both factors range from 1-5. The system shall classify risks as Low, '
                     'Medium, High, or Critical based on the calculated rating.', style='List Bullet 2')
    
    doc.add_paragraph('FR5: Report Generation', style='List Bullet')
    doc.add_paragraph('   The system shall generate comprehensive HSE reports in DOCX format containing: '
                     'Project Information, Inspection Summary, Detected Hazards, Risk Analysis, Control '
                     'Measures, Recommendations, Responsible Parties, and Follow-up Verification. The system '
                     'shall also generate JSON and raw text outputs.', style='List Bullet 2')
    
    doc.add_paragraph('FR6: Metadata Collection', style='List Bullet')
    doc.add_paragraph('   The system shall collect and store project name, site location, inspection '
                     'date/time (Asia/Riyadh timezone), inspector name, and verification details.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('FR7: Control Measures Generation', style='List Bullet')
    doc.add_paragraph('   The system shall provide control measures following the hierarchy of controls: '
                     'Elimination, Substitution, Engineering Controls, Administrative Controls, and PPE.', 
                     style='List Bullet 2')
    
    doc.add_heading('2.3 Non-Functional Requirements', level=2)
    
    doc.add_paragraph('NFR1: Performance', style='List Bullet')
    doc.add_paragraph('   The system shall process a single image analysis within 30-60 seconds, '
                     'including YOLO detection and vision-language model analysis.', style='List Bullet 2')
    
    doc.add_paragraph('NFR2: Reliability', style='List Bullet')
    doc.add_paragraph('   The system shall handle API failures gracefully and provide appropriate error '
                     'messages to users.', style='List Bullet 2')
    
    doc.add_paragraph('NFR3: Usability', style='List Bullet')
    doc.add_paragraph('   The web interface shall be intuitive and require minimal training. Users '
                     'should be able to complete an inspection upload in less than 2 minutes.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('NFR4: Security', style='List Bullet')
    doc.add_paragraph('   The system shall securely handle uploaded images and not store sensitive '
                     'data longer than necessary. API keys shall be stored as environment variables.', 
                     style='List Bullet 2')
    
    doc.add_paragraph('NFR5: Compatibility', style='List Bullet')
    doc.add_paragraph('   The system shall work on modern web browsers (Chrome, Firefox, Safari, Edge) '
                     'and support responsive design for mobile devices.', style='List Bullet 2')
    
    doc.add_paragraph('NFR6: Compliance', style='List Bullet')
    doc.add_paragraph('   Generated reports shall comply with Middle-East HSE inspection standards '
                     'and formats required by regulatory bodies in Saudi Arabia.', style='List Bullet 2')
    
    doc.add_heading('2.4 Usability Requirements', level=2)
    
    doc.add_paragraph('UR1: The user interface shall provide clear visual feedback during processing.', 
                     style='List Bullet')
    doc.add_paragraph('UR2: Error messages shall be user-friendly and suggest corrective actions.', 
                     style='List Bullet')
    doc.add_paragraph('UR3: The system shall provide downloadable reports in multiple formats (DOCX, JSON).', 
                     style='List Bullet')
    doc.add_paragraph('UR4: Results shall be displayed in an organized dashboard format showing detections, '
                     'risk levels, and key findings.', style='List Bullet')
    
    doc.add_page_break()
    
    # 3. SYSTEM MODELING
    heading3 = doc.add_heading('3. System Modeling', level=1)
    
    doc.add_heading('3.1 Use Cases', level=2)
    
    doc.add_paragraph('The AUHSE system has the following primary use cases:')
    
    doc.add_heading('Use Case 1: Upload and Analyze Site Photo', level=3)
    doc.add_paragraph('Actor: HSE Inspector')
    doc.add_paragraph('Preconditions: Inspector has access to the web application and a site photo')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('   1. Inspector navigates to the web application', style='List Number')
    doc.add_paragraph('   2. Inspector enters project name, site location, inspection by, and verified by', 
                     style='List Number')
    doc.add_paragraph('   3. Inspector uploads a site photo', style='List Number')
    doc.add_paragraph('   4. System validates image format and size', style='List Number')
    doc.add_paragraph('   5. System runs YOLO detection to check relevance', style='List Number')
    doc.add_paragraph('   6. If relevant, system processes image with vision-language model', 
                     style='List Number')
    doc.add_paragraph('   7. System generates and displays HSE report', style='List Number')
    doc.add_paragraph('   8. Inspector reviews results and downloads reports', style='List Number')
    
    doc.add_paragraph('Postconditions: HSE report is generated and available for download')
    doc.add_paragraph('Alternative Flow: If image is not relevant, system displays appropriate message')
    
    doc.add_heading('Use Case 2: Review Detection Results', level=3)
    doc.add_paragraph('Actor: HSE Inspector')
    doc.add_paragraph('Description: Inspector reviews the detected objects, hazards, and risk assessments')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('   1. After analysis, system displays detection results', style='List Number')
    doc.add_paragraph('   2. Inspector views detected objects with confidence scores', style='List Number')
    doc.add_paragraph('   3. Inspector reviews identified hazards and locations', style='List Number')
    doc.add_paragraph('   4. Inspector reviews risk ratings and control measures', style='List Number')
    
    doc.add_heading('Use Case 3: Download Reports', level=3)
    doc.add_paragraph('Actor: HSE Inspector')
    doc.add_paragraph('Description: Inspector downloads generated reports in various formats')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('   1. After report generation, system provides download links', style='List Number')
    doc.add_paragraph('   2. Inspector clicks on desired format (DOCX, JSON, or raw text)', style='List Number')
    doc.add_paragraph('   3. System serves the requested file for download', style='List Number')
    
    doc.add_heading('3.2 Class Diagram', level=2)
    
    doc.add_paragraph('The system follows an object-oriented design with the following main classes:')
    
    doc.add_paragraph('FlaskApp: Main application class managing routes and configuration', style='List Bullet')
    doc.add_paragraph('   - Attributes: config, upload_folder, output_folder, yolo_model_path', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Methods: create_app(), register_routes(), index(), analyze()', 
                     style='List Bullet 2')
    
    doc.add_paragraph('HSEPipeline: Processes HSE inspection requests', style='List Bullet')
    doc.add_paragraph('   - Attributes: config, image_path, project_name, site_location', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Methods: process_hse_request(), run_yolo_gate()', 
                     style='List Bullet 2')
    
    doc.add_paragraph('YoloDetection: Represents a single object detection', style='List Bullet')
    doc.add_paragraph('   - Attributes: class_name, confidence, bbox', style='List Bullet 2')
    
    doc.add_paragraph('HSEReportGenerator: Generates HSE reports using vision-language model', style='List Bullet')
    doc.add_paragraph('   - Attributes: image_path, project_info, api_endpoint', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Methods: generate_hse_report(), build_hse_prompt(), call_openrouter_mm(), '
                     'write_hse_docx()', style='List Bullet 2')
    
    doc.add_paragraph('HSEReport: Structured data class for HSE report', style='List Bullet')
    doc.add_paragraph('   - Attributes: project_information, inspection_summary, detected_hazards, '
                     'risk_analysis, control_measures, recommendations', style='List Bullet 2')
    
    doc.add_paragraph(
        'Relationships: FlaskApp uses HSEPipeline, HSEPipeline uses YoloDetection and HSEReportGenerator, '
        'HSEReportGenerator creates HSEReport objects.'
    )
    
    doc.add_heading('3.3 Sequence Diagrams', level=2)
    
    doc.add_paragraph('Sequence Diagram: Complete Inspection Flow')
    doc.add_paragraph(
        '1. User → FlaskApp: Submit form with project details and image'
    )
    doc.add_paragraph(
        '2. FlaskApp → HSEPipeline: process_hse_request(image_path, project_info)'
    )
    doc.add_paragraph(
        '3. HSEPipeline → YOLO Model: run_yolo_gate(image_path)'
    )
    doc.add_paragraph(
        '4. YOLO Model → HSEPipeline: Return detections and relevance flag'
    )
    doc.add_paragraph(
        '5. HSEPipeline → HSEReportGenerator: generate_hse_report(image_path, metadata)'
    )
    doc.add_paragraph(
        '6. HSEReportGenerator → OpenRouter API: call_openrouter_mm(prompt, image)'
    )
    doc.add_paragraph(
        '7. OpenRouter API → InternVL3: Process vision-language request'
    )
    doc.add_paragraph(
        '8. InternVL3 → OpenRouter API: Return structured JSON report'
    )
    doc.add_paragraph(
        '9. OpenRouter API → HSEReportGenerator: Return JSON response'
    )
    doc.add_paragraph(
        '10. HSEReportGenerator → FileSystem: write_hse_docx(report_data)'
    )
    doc.add_paragraph(
        '11. HSEReportGenerator → HSEPipeline: Return report paths and data'
    )
    doc.add_paragraph(
        '12. HSEPipeline → FlaskApp: Return result dictionary'
    )
    doc.add_paragraph(
        '13. FlaskApp → User: Render results page with download links'
    )
    
    doc.add_page_break()
    
    # 4. DATABASE DESIGN
    heading4 = doc.add_heading('4. Database Design', level=1)
    
    doc.add_heading('4.1 Introduction', level=2)
    doc.add_paragraph(
        'The AUHSE system currently operates with a file-based storage approach. However, for production '
        'deployment, a database design is presented to support scalability, data persistence, and advanced '
        'features such as report history and analytics.'
    )
    
    doc.add_heading('4.2 Data Requirements', level=2)
    doc.add_paragraph('The system needs to store the following data:')
    
    doc.add_paragraph('• Project Information: Project names, site locations, dates', style='List Bullet')
    doc.add_paragraph('• Inspection Records: Each inspection session with metadata', style='List Bullet')
    doc.add_paragraph('• Detection Results: YOLO detections with confidence scores and bounding boxes', 
                     style='List Bullet')
    doc.add_paragraph('• HSE Reports: Structured hazard data, risk assessments, control measures', 
                     style='List Bullet')
    doc.add_paragraph('• User Information: Inspector names, verification details', style='List Bullet')
    doc.add_paragraph('• File References: Paths to uploaded images and generated reports', style='List Bullet')
    
    doc.add_heading('4.3 Entity Relationship Model', level=2)
    
    doc.add_paragraph('Primary Entities:')
    
    doc.add_paragraph('Inspections Table (Primary Key: inspection_id)', style='List Bullet')
    doc.add_paragraph('   - inspection_id (INT, PRIMARY KEY, AUTO_INCREMENT)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - project_name (VARCHAR(255), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - site_location (VARCHAR(255), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - inspection_date (DATETIME, NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - inspection_by (VARCHAR(255), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - verified_by (VARCHAR(255), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - image_path (VARCHAR(500), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - output_folder (VARCHAR(500), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - status (ENUM("processing", "completed", "failed"), NOT NULL)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - created_at (TIMESTAMP, DEFAULT CURRENT_TIMESTAMP)', 
                     style='List Bullet 2')
    
    doc.add_paragraph('Detections Table (Primary Key: detection_id)', style='List Bullet')
    doc.add_paragraph('   - detection_id (INT, PRIMARY KEY, AUTO_INCREMENT)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - inspection_id (INT, FOREIGN KEY → Inspections)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - class_name (VARCHAR(100), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - confidence (DECIMAL(5,3), NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - bbox_x1, bbox_y1, bbox_x2, bbox_y2 (DECIMAL(10,2))', 
                     style='List Bullet 2')
    
    doc.add_paragraph('Hazards Table (Primary Key: hazard_id)', style='List Bullet')
    doc.add_paragraph('   - hazard_id (INT, PRIMARY KEY, AUTO_INCREMENT)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - inspection_id (INT, FOREIGN KEY → Inspections)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - image_reference (INT)', style='List Bullet 2')
    doc.add_paragraph('   - description (TEXT, NOT NULL)', style='List Bullet 2')
    doc.add_paragraph('   - causes (TEXT)', style='List Bullet 2')
    doc.add_paragraph('   - location_on_site (VARCHAR(255))', style='List Bullet 2')
    
    doc.add_paragraph('Risk Analysis Table (Primary Key: risk_id)', style='List Bullet')
    doc.add_paragraph('   - risk_id (INT, PRIMARY KEY, AUTO_INCREMENT)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - inspection_id (INT, FOREIGN KEY → Inspections, UNIQUE)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - likelihood (INT, CHECK 1-5)', style='List Bullet 2')
    doc.add_paragraph('   - severity (INT, CHECK 1-5)', style='List Bullet 2')
    doc.add_paragraph('   - risk_rating (INT, CHECK 1-25)', style='List Bullet 2')
    doc.add_paragraph('   - risk_level (ENUM("Low", "Medium", "High", "Critical"))', 
                     style='List Bullet 2')
    
    doc.add_paragraph('Control Measures Table (Primary Key: control_id)', style='List Bullet')
    doc.add_paragraph('   - control_id (INT, PRIMARY KEY, AUTO_INCREMENT)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - inspection_id (INT, FOREIGN KEY → Inspections)', 
                     style='List Bullet 2')
    doc.add_paragraph('   - measure_type (ENUM("elimination", "substitution", "engineering", '
                     '"administrative", "ppe"))', style='List Bullet 2')
    doc.add_paragraph('   - description (TEXT, NOT NULL)', style='List Bullet 2')
    
    doc.add_heading('4.4 Normalization', level=2)
    doc.add_paragraph(
        'The database design follows Third Normal Form (3NF):'
    )
    doc.add_paragraph('• First Normal Form: All attributes contain atomic values, no repeating groups', 
                     style='List Bullet')
    doc.add_paragraph('• Second Normal Form: All non-key attributes are fully dependent on primary keys', 
                     style='List Bullet')
    doc.add_paragraph('• Third Normal Form: No transitive dependencies; all attributes depend only on the primary key', 
                     style='List Bullet')
    
    doc.add_paragraph(
        'Foreign key relationships ensure referential integrity. The design minimizes data redundancy while '
        'maintaining query efficiency through appropriate indexing.'
    )
    
    doc.add_heading('4.5 Indexing Strategy', level=2)
    doc.add_paragraph('Recommended indexes for performance optimization:')
    doc.add_paragraph('• Inspections: INDEX on inspection_date, INDEX on project_name', 
                     style='List Bullet')
    doc.add_paragraph('• Detections: INDEX on inspection_id, INDEX on class_name', 
                     style='List Bullet')
    doc.add_paragraph('• Hazards: INDEX on inspection_id', style='List Bullet')
    doc.add_paragraph('• Risk Analysis: INDEX on inspection_id (unique), INDEX on risk_level', 
                     style='List Bullet')
    
    doc.add_page_break()
    
    # 5. ARCHITECTURAL DESIGN
    heading5 = doc.add_heading('5. Architectural Design', level=1)
    
    doc.add_heading('5.1 Introduction', level=2)
    doc.add_paragraph(
        'The AUHSE system follows a three-tier architecture pattern, separating the presentation layer, '
        'business logic layer, and data access layer. This architecture ensures modularity, scalability, '
        'and maintainability.'
    )
    
    doc.add_heading('5.2 Three-Tier Architecture', level=2)
    
    doc.add_heading('5.2.1 Presentation Layer', level=3)
    doc.add_paragraph(
        'The presentation layer is implemented using Flask web framework with Jinja2 templating. '
        'This layer handles:'
    )
    doc.add_paragraph('• User interface rendering (HTML/CSS)', style='List Bullet')
    doc.add_paragraph('• HTTP request/response handling', style='List Bullet')
    doc.add_paragraph('• Form validation and user input processing', style='List Bullet')
    doc.add_paragraph('• File upload handling', style='List Bullet')
    doc.add_paragraph('• Result display and report download links', style='List Bullet')
    
    doc.add_paragraph(
        'Key components: index.html template, Flask routes (/, /analyze, /reports/<folder>/<asset>), '
        'and static assets (CSS, images).'
    )
    
    doc.add_heading('5.2.2 Business Logic Layer', level=3)
    doc.add_paragraph('The business logic layer contains the core processing components:')
    
    doc.add_paragraph('Pipeline Module (pipeline.py):', style='List Bullet')
    doc.add_paragraph('   - Orchestrates the complete inspection workflow', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Manages YOLO-based relevance detection', style='List Bullet 2')
    doc.add_paragraph('   - Coordinates report generation process', style='List Bullet 2')
    
    doc.add_paragraph('HSE Report Generator Module (LLM_VLM.py):', style='List Bullet')
    doc.add_paragraph('   - Interfaces with OpenRouter API for vision-language model access', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Constructs structured prompts for hazard analysis', style='List Bullet 2')
    doc.add_paragraph('   - Processes JSON responses and generates DOCX reports', 
                     style='List Bullet 2')
    
    doc.add_paragraph(
        'This layer encapsulates all business rules, including risk calculation algorithms, report '
        'formatting standards, and API integration logic.'
    )
    
    doc.add_heading('5.2.3 Data Access Layer', level=3)
    doc.add_paragraph('Currently implemented as file-based storage:')
    doc.add_paragraph('• Upload folder: Stores uploaded images temporarily', style='List Bullet')
    doc.add_paragraph('• Output folder: Stores generated reports organized by timestamp', 
                     style='List Bullet')
    doc.add_paragraph('• File system operations: Managed through Python Path and shutil modules', 
                     style='List Bullet')
    
    doc.add_paragraph(
        'For production deployment, this layer would interface with a relational database (MySQL/PostgreSQL) '
        'or cloud storage services (AWS S3, Azure Blob Storage) for persistent data storage.'
    )
    
    doc.add_heading('5.3 Component Architecture', level=2)
    doc.add_paragraph('The system consists of the following key components:')
    
    doc.add_paragraph('1. Flask Application (app.py)', style='List Number')
    doc.add_paragraph('   - Entry point and route handlers', style='List Bullet 2')
    doc.add_paragraph('   - Configuration management', style='List Bullet 2')
    doc.add_paragraph('   - Error handling and flash messages', style='List Bullet 2')
    
    doc.add_paragraph('2. Processing Pipeline (pipeline.py)', style='List Number')
    doc.add_paragraph('   - YOLO model integration and caching', style='List Bullet 2')
    doc.add_paragraph('   - Relevance detection logic', style='List Bullet 2')
    doc.add_paragraph('   - Request orchestration', style='List Bullet 2')
    
    doc.add_paragraph('3. AI Integration (LLM_VLM.py)', style='List Number')
    doc.add_paragraph('   - OpenRouter API client', style='List Bullet 2')
    doc.add_paragraph('   - Prompt engineering', style='List Bullet 2')
    doc.add_paragraph('   - JSON parsing and validation', style='List Bullet 2')
    doc.add_paragraph('   - DOCX report generation', style='List Bullet 2')
    
    doc.add_paragraph('4. YOLO Model (best.pt)', style='List Number')
    doc.add_paragraph('   - Pre-trained object detection model', style='List Bullet 2')
    doc.add_paragraph('   - Loaded via ultralytics library', style='List Bullet 2')
    
    doc.add_heading('5.4 API Integration', level=2)
    doc.add_paragraph(
        'The system integrates with external APIs for AI capabilities:'
    )
    doc.add_paragraph('• OpenRouter API: Provides access to InternVL3-78B vision-language model', 
                     style='List Bullet')
    doc.add_paragraph('• Communication: RESTful HTTP POST requests with JSON payloads', 
                     style='List Bullet')
    doc.add_paragraph('• Authentication: Bearer token via API key stored in environment variables', 
                     style='List Bullet')
    doc.add_paragraph('• Image encoding: Base64-encoded images embedded in API requests', 
                     style='List Bullet')
    
    doc.add_heading('5.5 Benefits of Three-Tier Architecture', level=2)
    doc.add_paragraph('• Separation of Concerns: Each layer has distinct responsibilities', 
                     style='List Bullet')
    doc.add_paragraph('• Scalability: Components can be scaled independently', style='List Bullet')
    doc.add_paragraph('• Maintainability: Changes to one layer do not affect others', style='List Bullet')
    doc.add_paragraph('• Testability: Each layer can be tested in isolation', style='List Bullet')
    doc.add_paragraph('• Reusability: Business logic can be reused across different interfaces', 
                     style='List Bullet')
    doc.add_paragraph('• Flexibility: Easy to replace or upgrade individual components', 
                     style='List Bullet')
    
    doc.add_page_break()
    
    # 6. USER INTERFACE DESIGN
    heading6 = doc.add_heading('6. User Interface Design', level=1)
    
    doc.add_heading('6.1 Introduction', level=2)
    doc.add_paragraph(
        'The user interface design focuses on simplicity, efficiency, and user experience. The interface '
        'follows modern web design principles to ensure intuitive navigation and clear visual feedback.'
    )
    
    doc.add_heading('6.2 Screen Design', level=2)
    
    doc.add_heading('6.2.1 Main Page (Index)', level=3)
    doc.add_paragraph('Layout: Split-panel design with form on the left and informational content on the right')
    
    doc.add_paragraph('Left Panel - Inspection Form:', style='List Bullet')
    doc.add_paragraph('   - Header: "Launch a site inspection"', style='List Bullet 2')
    doc.add_paragraph('   - Form fields:', style='List Bullet 2')
    doc.add_paragraph('     • Project name (text input, required)', style='List Bullet 3')
    doc.add_paragraph('     • Site location (text input, required)', style='List Bullet 3')
    doc.add_paragraph('     • Inspection by (text input, required)', style='List Bullet 3')
    doc.add_paragraph('     • Verified by (text input, required)', style='List Bullet 3')
    doc.add_paragraph('   - File upload area: Drag-and-drop zone with visual indicators', 
                     style='List Bullet 2')
    doc.add_paragraph('   - Submit button: "Generate HSE intelligence"', style='List Bullet 2')
    
    doc.add_paragraph('Right Panel - Information Display:', style='List Bullet')
    doc.add_paragraph('   - Hero image: Construction site background', style='List Bullet 2')
    doc.add_paragraph('   - System description and key features', style='List Bullet 2')
    doc.add_paragraph('   - Statistics panel: System capabilities and metrics', style='List Bullet 2')
    
    doc.add_heading('6.2.2 Results Page', level=3)
    doc.add_paragraph('Displayed after successful analysis:')
    
    doc.add_paragraph('Header Section:', style='List Bullet')
    doc.add_paragraph('   - Inspection Dashboard title', style='List Bullet 2')
    doc.add_paragraph('   - Generation timestamp', style='List Bullet 2')
    doc.add_paragraph('   - Risk level badge (color-coded: Low/Medium/High/Critical)', 
                     style='List Bullet 2')
    
    doc.add_paragraph('Results Body (Two-column layout):', style='List Bullet')
    doc.add_paragraph('   Left Column:', style='List Bullet 2')
    doc.add_paragraph('     • Detections list with class names and confidence percentages', 
                     style='List Bullet 3')
    doc.add_paragraph('     • Risk analysis showing L×S calculation', style='List Bullet 3')
    
    doc.add_paragraph('   Right Column:', style='List Bullet 2')
    doc.add_paragraph('     • Detected hazards with descriptions, causes, and locations', 
                     style='List Bullet 3')
    doc.add_paragraph('     • Recommended control measures (hierarchy of controls)', 
                     style='List Bullet 3')
    
    doc.add_paragraph('Footer Section:', style='List Bullet')
    doc.add_paragraph('   - Download buttons: DOCX Report, Raw JSON, Model Output', 
                     style='List Bullet 2')
    
    doc.add_heading('6.3 Design Principles Applied', level=2)
    
    doc.add_heading('6.3.1 Visibility of System Status', level=3)
    doc.add_paragraph(
        'The system provides clear feedback at every stage: form validation messages, processing indicators, '
        'and result status displays. Risk levels are prominently displayed with color-coded badges.'
    )
    
    doc.add_heading('6.3.2 Match Between System and Real World', level=3)
    doc.add_paragraph(
        'The interface uses construction industry terminology (project name, site location, inspection by) '
        'that matches users\' familiar workflow. The report format follows standard HSE inspection templates.'
    )
    
    doc.add_heading('6.3.3 User Control and Freedom', level=3)
    doc.add_paragraph(
        'Users can easily correct inputs, re-upload images, and download reports in multiple formats. '
        'The interface supports standard browser navigation (back/forward buttons).'
    )
    
    doc.add_heading('6.3.4 Consistency and Standards', level=3)
    doc.add_paragraph(
        'Consistent button styles, form layouts, and color schemes throughout the application. Follows '
        'standard web conventions (form labels, required field indicators, error messages).'
    )
    
    doc.add_heading('6.3.5 Error Prevention', level=3)
    doc.add_paragraph(
        'Form validation prevents submission with missing required fields. File type and size validation '
        'prevents invalid uploads. Clear error messages guide users to correct issues.'
    )
    
    doc.add_heading('6.3.6 Recognition Rather Than Recall', level=3)
    doc.add_paragraph(
        'All necessary information is visible on the screen. Form fields retain values after submission. '
        'Download links are clearly labeled with file types.'
    )
    
    doc.add_heading('6.3.7 Flexibility and Efficiency of Use', level=3)
    doc.add_paragraph(
        'The interface accommodates both novice and experienced users. Drag-and-drop file upload '
        'provides quick access for frequent users while maintaining a browse option for others.'
    )
    
    doc.add_heading('6.3.8 Aesthetic and Minimalist Design', level=3)
    doc.add_paragraph(
        'Clean, uncluttered interface focuses on essential functionality. Gradient backgrounds and '
        'subtle shadows provide visual depth without distraction.'
    )
    
    doc.add_heading('6.3.9 Error Recovery', level=3)
    doc.add_paragraph(
        'Clear, actionable error messages guide users. For API failures or processing errors, the system '
        'provides specific information about what went wrong and suggests next steps.'
    )
    
    doc.add_heading('6.3.10 Help and Documentation', level=3)
    doc.add_paragraph(
        'Inline help text in form fields (placeholders) and descriptive subtitles explain system '
        'functionality. The information panel provides context about system capabilities.'
    )
    
    doc.add_heading('6.4 Responsive Design', level=2)
    doc.add_paragraph(
        'The interface uses CSS media queries to adapt to different screen sizes. On mobile devices, '
        'the two-column layout converts to a single-column stack for optimal viewing. Touch-friendly '
        'button sizes and spacing ensure usability on tablets and smartphones.'
    )
    
    doc.add_heading('6.5 Color Scheme and Typography', level=2)
    doc.add_paragraph('• Color Scheme: Professional construction-themed palette with gradient backgrounds', 
                     style='List Bullet')
    doc.add_paragraph('• Risk Level Colors: Low (green), Medium (yellow), High (orange), Critical (red)', 
                     style='List Bullet')
    doc.add_paragraph('• Typography: Modern sans-serif fonts for headings, readable serif fonts for body text', 
                     style='List Bullet')
    doc.add_paragraph('• Contrast: High contrast ratios ensure accessibility and readability', 
                     style='List Bullet')
    
    doc.add_page_break()
    
    # 7. FUTURE WORK
    heading7 = doc.add_heading('7. Future Work', level=1)
    
    doc.add_heading('7.1 Enhanced AI Capabilities', level=2)
    doc.add_paragraph('• Multi-image analysis: Support for multiple images per inspection to provide '
                     'comprehensive site coverage', style='List Bullet')
    doc.add_paragraph('• Video analysis: Real-time processing of construction site video feeds for '
                     'continuous monitoring', style='List Bullet')
    doc.add_paragraph('• Predictive analytics: Machine learning models to predict potential hazards '
                     'based on historical data', style='List Bullet')
    doc.add_paragraph('• Custom model training: Fine-tune YOLO model on construction site-specific '
                     'datasets for improved accuracy', style='List Bullet')
    
    doc.add_heading('7.2 Database Integration', level=2)
    doc.add_paragraph('• Implement full database schema with MySQL or PostgreSQL', style='List Bullet')
    doc.add_paragraph('• User authentication and authorization system', style='List Bullet')
    doc.add_paragraph('• Report history and search functionality', style='List Bullet')
    doc.add_paragraph('• Analytics dashboard with trend analysis and statistics', style='List Bullet')
    doc.add_paragraph('• Data export capabilities for compliance reporting', style='List Bullet')
    
    doc.add_heading('7.3 Advanced Features', level=2)
    doc.add_paragraph('• Mobile application: Native iOS and Android apps for on-site inspections', 
                     style='List Bullet')
    doc.add_paragraph('• Integration with construction management software (Procore, PlanGrid, etc.)', 
                     style='List Bullet')
    doc.add_paragraph('• Real-time notifications: Alert system for critical risk detections', 
                     style='List Bullet')
    doc.add_paragraph('• Multi-language support: Arabic and English interface and reports', 
                     style='List Bullet')
    doc.add_paragraph('• Batch processing: Analyze multiple inspections simultaneously', 
                     style='List Bullet')
    doc.add_paragraph('• API for third-party integrations', style='List Bullet')
    
    doc.add_heading('7.4 Performance Optimization', level=2)
    doc.add_paragraph('• Implement caching mechanisms for frequently accessed data', style='List Bullet')
    doc.add_paragraph('• Async processing: Queue-based system for handling multiple concurrent requests', 
                     style='List Bullet')
    doc.add_paragraph('• CDN integration: Optimize static asset delivery', style='List Bullet')
    doc.add_paragraph('• Image optimization: Compression and thumbnail generation', style='List Bullet')
    doc.add_paragraph('• Database query optimization with proper indexing', style='List Bullet')
    
    doc.add_heading('7.5 Compliance and Standards', level=2)
    doc.add_paragraph('• Integration with Saudi Building Code (SBC) compliance standards', 
                     style='List Bullet')
    doc.add_paragraph('• Support for Civil Defense requirements', style='List Bullet')
    doc.add_paragraph('• HRSD OSH (Occupational Safety and Health) standard compliance', 
                     style='List Bullet')
    doc.add_paragraph('• ISO 45001 (Occupational Health and Safety) alignment', style='List Bullet')
    doc.add_paragraph('• Audit trail and compliance reporting features', style='List Bullet')
    
    doc.add_heading('7.6 User Experience Enhancements', level=2)
    doc.add_paragraph('• Interactive map integration for site location visualization', style='List Bullet')
    doc.add_paragraph('• Report customization: User-defined templates and fields', style='List Bullet')
    doc.add_paragraph('• Collaborative features: Multi-user inspection workflows', style='List Bullet')
    doc.add_paragraph('• Workflow automation: Automated follow-up and verification reminders', 
                     style='List Bullet')
    doc.add_paragraph('• Advanced filtering and sorting for inspection history', style='List Bullet')
    
    doc.add_page_break()
    
    # 8. REFERENCES
    heading8 = doc.add_heading('8. References', level=1)
    
    doc.add_paragraph('1. Ultralytics. (2023). YOLOv8 Documentation. https://docs.ultralytics.com/')
    
    doc.add_paragraph('2. OpenGVLab. (2024). InternVL: Scaling up Vision Foundation Models and Aligning for '
                     'Generic Visual-Linguistic Tasks. arXiv preprint.')
    
    doc.add_paragraph('3. Flask Development Team. (2023). Flask Web Framework Documentation. '
                     'https://flask.palletsprojects.com/')
    
    doc.add_paragraph('4. Saudi Building Code (SBC). Saudi Building Code National Committee. '
                     'https://sbc.gov.sa/')
    
    doc.add_paragraph('5. ISO 45001:2018. Occupational health and safety management systems — Requirements '
                     'with guidance for use. International Organization for Standardization.')
    
    doc.add_paragraph('6. Redmon, J., et al. (2016). You Only Look Once: Unified, Real-Time Object Detection. '
                     'Proceedings of the IEEE conference on computer vision and pattern recognition.')
    
    doc.add_paragraph('7. Pillow (PIL Fork) Documentation. (2023). Python Imaging Library. '
                     'https://pillow.readthedocs.io/')
    
    doc.add_paragraph('8. python-docx Documentation. (2023). Create and modify Microsoft Word .docx files. '
                     'https://python-docx.readthedocs.io/')
    
    doc.add_paragraph('9. OpenRouter. (2024). API Documentation. https://openrouter.ai/docs')
    
    doc.add_paragraph('10. Construction Industry Standards. (2023). Health, Safety, and Environment Guidelines '
                     'for Construction Sites in Saudi Arabia.')
    
    doc.add_paragraph()
    doc.add_paragraph(f'Document Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = '/Volumes/MahmoudHard/Downloads/gradution_project/Final_Presentation_AUHSE_System.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
